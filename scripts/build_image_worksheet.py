#!/usr/bin/env python3
"""Genera un documento Word para decidir a mano el tamaño de cada imagen.

Flujo de trabajo:
1. Ejecuta este script. Genera `image-worksheet/tamanos-imagenes.docx` y
   `image-worksheet/manifest.json`.
2. Abre el .docx, redimensiona (arrastrando una esquina, para mantener la
   proporción) las imágenes que quieras cambiar y deja las demás tal cual.
3. Exporta/guarda ese documento como PDF, sin cambiar la escala de impresión
   ("tamaño real").
4. Pásame el PDF resultante y ejecuto `apply_image_sizes.py` con él para
   aplicar esos tamaños a docs/*.md.

No toques el texto en gris bajo cada imagen (el identificador `IMGxxx`):
es lo que uso para volver a encontrar cada imagen en el PDF exportado.
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from image_inventory import build_inventory
from manual_nav import DOCS_DIR, load_nav, ordered_pages

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "image-worksheet"
DOCX_PATH = OUT_DIR / "tamanos-imagenes.docx"
MANIFEST_PATH = OUT_DIR / "manifest.json"

BOX_W_IN = 2.6
BOX_H_IN = 3.2
PX_PER_IN = 96


def page_titles():
    """título legible por fichero .md, tomado del primer <h1>/# del fichero."""
    titles = {}
    for f in DOCS_DIR.rglob("*.md"):
        rel = f.relative_to(DOCS_DIR).as_posix()
        for line in f.read_text(encoding="utf-8").splitlines():
            if line.startswith("# "):
                titles[rel] = line[2:].strip()
                break
        titles.setdefault(rel, rel)
    return titles


def fitted_width_in(file_path: Path) -> float:
    with Image.open(file_path) as img:
        native_w, native_h = img.size
    w_in = native_w / PX_PER_IN
    h_in = native_h / PX_PER_IN
    scale = min(BOX_W_IN / w_in, BOX_H_IN / h_in, 1.0)
    return w_in * scale


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p


def main():
    OUT_DIR.mkdir(exist_ok=True)
    nav, site_title = load_nav()
    pages = ordered_pages(nav, include_home=True)
    titles = page_titles()

    usages = build_inventory()
    by_page = {}
    for u in usages:
        by_page.setdefault(u.md_path, []).append(u)

    doc = Document()
    doc.add_heading(f"Tamaños de imágenes — {site_title}", level=0)
    intro = doc.add_paragraph()
    intro.add_run(
        "Redimensiona (arrastrando una esquina) las imágenes que quieras cambiar "
        "de tamaño; deja las demás tal cual. No edites el texto gris bajo cada "
        "imagen. Cuando termines, exporta este documento a PDF a tamaño real "
        "(sin ajustar a página) y devuélvemelo."
    ).italic = True

    manifest = []

    for _depth, md_path in pages:
        page_usages = by_page.get(md_path, [])
        if not page_usages:
            continue

        doc.add_heading(f"{titles.get(md_path, md_path)}", level=1)
        doc.add_paragraph(f"docs/{md_path}").italic = True

        for u in page_usages:
            width_in = fitted_width_in(u.file_path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True
            run = p.add_run()
            run.add_picture(str(u.file_path), width=Inches(width_in))

            with Image.open(u.file_path) as img:
                native_w, native_h = img.size
            current = f"{u.current_width} px" if u.current_width else "ancho automático"
            caption = f"{u.image_id} · {u.file_path.name} · nativo {native_w}×{native_h} px · actual: {current}"
            add_caption(doc, caption)

            manifest.append({
                "id": u.image_id,
                "md_path": u.md_path,
                "src": u.src,
                "kind": u.kind,
                "occurrence_index": u.occurrence_index,
                "current_width": u.current_width,
                "filename": u.file_path.name,
                "native_width": native_w,
                "native_height": native_h,
                "initial_width_in": round(width_in, 4),
            })

    doc.save(DOCX_PATH)
    MANIFEST_PATH.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Documento generado en {DOCX_PATH} ({len(manifest)} imágenes)")
    print(f"Manifiesto en {MANIFEST_PATH}")


if __name__ == "__main__":
    main()
